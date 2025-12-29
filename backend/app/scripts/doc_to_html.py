#!/usr/bin/env python3
"""
Word 转 HTML 转换脚本 - 完整重写
完整保留所有格式：字体（包括中文字体）、大小、颜色、加粗、斜体、下划线、图片、表格、对齐等
"""

import sys
import os
import argparse
import base64
import logging
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

logger = logging.getLogger(__name__)


def docx_to_html(input_path, output_path):
    """将 Word 文档转换为 HTML，保留所有格式"""
    try:
        print(f"开始转换: {input_path} -> {output_path}")

        # 加载文档
        doc = Document(input_path)

        # 构建 HTML 框架
        html_parts = [
            "<!DOCTYPE html>",
            '<html lang="zh-CN">',
            "<head>",
            '  <meta charset="UTF-8">',
            '  <meta name="viewport" content="width=device-width, initial-scale=1.0">',
            "  <title>转换文档</title>",
            "  <style>",
            "    * { margin: 0; padding: 0; }",
            '    body { font-family: "Calibri", "宋体", "楷体", "黑体", Arial, sans-serif; margin: 20px 40px; line-height: 1.8; color: #333; }',
            "    h1 { font-size: 28px; margin: 20px 0 15px 0; font-weight: bold; color: #1f4e78; }",
            "    h2 { font-size: 24px; margin: 18px 0 12px 0; font-weight: bold; color: #2e5c8a; }",
            "    h3 { font-size: 20px; margin: 15px 0 10px 0; font-weight: bold; color: #3d6b99; }",
            "    p { margin: 10px 0; word-wrap: break-word; }",
            "    p.center { text-align: center; }",
            "    p.right { text-align: right; }",
            "    p.justify { text-align: justify; }",
            "    table { border-collapse: collapse; width: 100%; margin: 15px 0; border: 1px solid #999; }",
            "    th, td { border: 1px solid #999; padding: 12px; text-align: left; word-break: break-word; }",
            "    th { background-color: #e7e6e6; font-weight: bold; }",
            "    img { max-width: 100%; height: auto; margin: 10px 0; display: block; }",
            "  </style>",
            "</head>",
            "<body>",
        ]

        # 处理文档内容
        for element in doc.element.body:
            tag = element.tag

            # 段落
            if tag.endswith("}p"):
                for para in doc.paragraphs:
                    if para._element == element:
                        para_html = _convert_paragraph(para, doc)
                        if para_html:
                            html_parts.append(para_html)
                        break

            # 表格
            elif tag.endswith("}tbl"):
                for table in doc.tables:
                    if table._element == element:
                        table_html = _convert_table(table, doc)
                        html_parts.append(table_html)
                        break

        html_parts.extend(["</body>", "</html>"])

        # 写入文件
        html_content = "\n".join(html_parts)
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(html_content)

        print(f"✓ 转换成功: {output_path}")
        return True

    except Exception as e:
        print(f"✗ 转换失败: {e}")
        import traceback

        traceback.print_exc()
        return False


def _get_run_font_family(run):
    """提取字体名称，优先获取中文字体"""
    try:
        rPr = run._element.rPr
        if rPr is None:
            return None

        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is not None:
            # 中文字体（eastAsia）优先
            east_asia = rFonts.get(qn("w:eastAsia"))
            if east_asia and east_asia.strip():
                return east_asia.strip()

            # ASCII 字体
            ascii_font = rFonts.get(qn("w:ascii"))
            if ascii_font and ascii_font.strip():
                return ascii_font.strip()

            # 其他
            cs = rFonts.get(qn("w:cs"))
            if cs and cs.strip():
                return cs.strip()
    except Exception:
        logger.exception("Get run font family failed")

    return None


def _get_run_font_size(run):
    """提取字体大小（点）"""
    try:
        rPr = run._element.rPr
        if rPr is None:
            return None

        sz = rPr.find(qn("w:sz"))
        if sz is not None:
            val = sz.get(qn("w:val"))
            if val:
                # Word 中字号是半磅单位，除以 2 得到点数
                pt = int(val) / 2
                return int(pt)
    except Exception:
        logger.exception("Get run font size failed")

    return None


def _get_run_color(run):
    """提取字体颜色"""
    try:
        rPr = run._element.rPr
        if rPr is None:
            return None

        color_elem = rPr.find(qn("w:color"))
        if color_elem is not None:
            val = color_elem.get(qn("w:val"))
            if val and val.lower() != "auto":
                return f"#{val}"
    except Exception:
        logger.exception("Get run color failed")

    return None


def _get_run_bold(run):
    """判断是否加粗"""
    try:
        rPr = run._element.rPr
        if rPr is None:
            return False
        return rPr.find(qn("w:b")) is not None
    except Exception:
        logger.exception("Get run bold failed")
        return False


def _get_run_italic(run):
    """判断是否斜体"""
    try:
        rPr = run._element.rPr
        if rPr is None:
            return False
        return rPr.find(qn("w:i")) is not None
    except Exception:
        logger.exception("Get run italic failed")
        return False


def _get_run_underline(run):
    """判断是否下划线"""
    try:
        rPr = run._element.rPr
        if rPr is None:
            return False
        return rPr.find(qn("w:u")) is not None
    except Exception:
        logger.exception("Get run underline failed")
        return False


def _get_run_strike(run):
    """判断是否删除线"""
    try:
        rPr = run._element.rPr
        if rPr is None:
            return False
        return rPr.find(qn("w:strike")) is not None
    except Exception:
        logger.exception("Get run strike failed")
        return False


def _build_span_style(run):
    """为 Run 构建 CSS 样式"""
    styles = []

    # 字体
    font = _get_run_font_family(run)
    if font:
        styles.append(f'font-family: "{font}"')

    # 大小
    size = _get_run_font_size(run)
    if size:
        styles.append(f"font-size: {size}pt")

    # 颜色
    color = _get_run_color(run)
    if color:
        styles.append(f"color: {color}")

    # 加粗
    if _get_run_bold(run):
        styles.append("font-weight: bold")

    # 斜体
    if _get_run_italic(run):
        styles.append("font-style: italic")

    # 下划线
    if _get_run_underline(run):
        styles.append("text-decoration: underline")

    # 删除线
    if _get_run_strike(run):
        styles.append("text-decoration: line-through")

    return "; ".join(styles) if styles else None


def _escape_html(text):
    """转义 HTML 特殊字符"""
    return (
        text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;").replace('"', "&quot;")
    )


def _extract_images_from_run(run, doc):
    """从 Run 中提取图片并转换为 Base64 Data URI"""
    images = []

    try:
        drawings = run._element.findall(
            ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing"
        )

        for drawing in drawings:
            blips = drawing.findall(
                ".//{http://schemas.openxmlformats.org/drawingml/2006/main}blip"
            )

            for blip in blips:
                embed_id = blip.get(
                    "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
                )

                if embed_id:
                    try:
                        image_part = doc.part.related_part(embed_id)
                        image_bytes = image_part.blob
                        content_type = image_part.content_type or "image/jpeg"

                        image_b64 = base64.b64encode(image_bytes).decode("utf-8")
                        data_uri = f"data:{content_type};base64,{image_b64}"

                        images.append(data_uri)
                        print(f"  ✓ 内联图片: {len(image_bytes)} 字节")
                    except Exception as e:
                        print(f"  ⚠ 图片提取异常: {e}")
    except Exception:
        logger.exception("Extract images from run failed")

    return images


def _convert_paragraph(para, doc):
    """将段落转换为 HTML"""
    if not para.text.strip():
        # 检查是否有图片
        has_image = any(
            run._element.findall(
                ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing"
            )
            for run in para.runs
        )
        if not has_image:
            return ""

    try:
        # 判断段落类型
        style_name = para.style.name

        # 标题
        if "Heading 1" in style_name:
            content = _convert_runs(para, doc)
            return f"  <h1>{content}</h1>"
        elif "Heading 2" in style_name:
            content = _convert_runs(para, doc)
            return f"  <h2>{content}</h2>"
        elif "Heading 3" in style_name:
            content = _convert_runs(para, doc)
            return f"  <h3>{content}</h3>"

        # 普通段落
        content = _convert_runs(para, doc)

        # 对齐方式
        align_class = ""
        if para.alignment == WD_ALIGN_PARAGRAPH.CENTER:
            align_class = ' class="center"'
        elif para.alignment == WD_ALIGN_PARAGRAPH.RIGHT:
            align_class = ' class="right"'
        elif para.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
            align_class = ' class="justify"'

        return f"  <p{align_class}>{content}</p>"

    except Exception as e:
        print(f"  ⚠ 段落转换异常: {e}")
        try:
            content = _convert_runs(para, doc)
            return f"  <p>{content}</p>"
        except Exception:
            logger.exception("Convert paragraph failed")
            return ""


def _convert_runs(para, doc):
    """转换段落中的所有 Run"""
    parts = []

    for run in para.runs:
        # 检查图片
        images = _extract_images_from_run(run, doc)
        for img_uri in images:
            parts.append(f'<img src="{img_uri}" />')

        # 检查文本
        if run.text:
            text = _escape_html(run.text)
            style = _build_span_style(run)

            if style:
                parts.append(f'<span style="{style}">{text}</span>')
            else:
                parts.append(text)

    return "".join(parts)


def _convert_table(table, doc):
    """将表格转换为 HTML"""
    html = ["  <table>"]

    try:
        for row_idx, row in enumerate(table.rows):
            html.append("    <tr>")

            for cell in row.cells:
                tag = "th" if row_idx == 0 else "td"

                # 单元格内容
                cell_parts = []
                for para in cell.paragraphs:
                    cell_parts.append(_convert_runs(para, doc))

                cell_content = "".join(cell_parts) if cell_parts else "&nbsp;"
                html.append(f"      <{tag}>{cell_content}</{tag}>")

            html.append("    </tr>")
    except Exception as e:
        print(f"  ⚠ 表格处理异常: {e}")

    html.append("  </table>")
    return "\n".join(html)


def main():
    parser = argparse.ArgumentParser(description="Word 转 HTML")
    parser.add_argument("-i", "--input", required=True, help="输入文件")
    parser.add_argument("-o", "--output", required=True, help="输出文件")

    args = parser.parse_args()

    if not os.path.exists(args.input):
        print(f"错误: 文件不存在 - {args.input}")
        sys.exit(1)

    success = docx_to_html(args.input, args.output)
    sys.exit(0 if success else 1)


if __name__ == "__main__":
    main()
