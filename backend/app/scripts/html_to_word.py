#!/usr/bin/env python3
"""
HTML 转 Word 转换脚本
基于 Apache POI 的 Java 实现思路，使用 python-docx 实现
"""

import sys
import os
import argparse
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from bs4 import BeautifulSoup
import re


def html_to_docx(input_path, output_path):
    """
    将 HTML 转换为 Word 文档
    类似 Java Htm2WordUtil 的实现
    """
    try:
        print(f"开始转换: {input_path} -> {output_path}")

        # 读取 HTML 内容
        with open(input_path, "r", encoding="utf-8") as f:
            html_content = f.read()

        # 使用 BeautifulSoup 解析 HTML
        soup = BeautifulSoup(html_content, "html.parser")

        # 创建 Word 文档
        doc = Document()

        # 设置文档属性
        set_document_properties(doc)

        # 处理 HTML 内容
        process_html_content(soup, doc)

        # 保存文档
        doc.save(output_path)

        # 验证输出文件
        if os.path.exists(output_path):
            file_size = os.path.getsize(output_path)
            print(f"转换成功: {output_path}")
            print(f"输出文件大小: {file_size} 字节")
            return True
        else:
            print("错误: 输出文件未生成")
            return False

    except Exception as e:
        print(f"HTML 转 Word 失败: {str(e)}")
        import traceback

        traceback.print_exc()
        return False


def set_document_properties(doc):
    """设置文档属性"""
    # 设置默认字体
    set_default_font(doc)


def set_default_font(doc):
    """设置默认字体（类似 Java 版本中的字体设置）"""
    style = doc.styles["Normal"]
    font = style.font
    font.name = "宋体"
    font.size = Pt(12)

    # 设置中文字体
    try:
        font.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    except Exception:
        pass


def process_html_content(soup, doc):
    """处理 HTML 内容"""
    # 获取 body 或整个文档
    body = soup.find("body") or soup

    # 处理标题
    process_title(soup, doc)

    # 处理正文内容
    process_body_elements(body, doc)


def process_title(soup, doc):
    """处理文档标题"""
    title = soup.find("title")
    if title:
        title_text = title.get_text().strip()
        if title_text:
            # 添加主标题
            heading = doc.add_heading(level=1)
            run = heading.add_run(title_text)
            run.font.name = "黑体"
            run.font.size = Pt(16)
            run.font.bold = True
            try:
                run.element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
            except Exception:
                pass
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # 添加空行
            doc.add_paragraph()


def process_body_elements(body, doc):
    """处理 body 中的元素"""
    for element in body.find_all(recursive=False):
        process_element(element, doc)


def process_element(element, doc):
    """处理单个 HTML 元素"""
    tag_name = element.name.lower()

    if tag_name in ["h1", "h2", "h3", "h4", "h5", "h6"]:
        process_heading(element, doc, int(tag_name[1]))

    elif tag_name == "p":
        process_paragraph(element, doc)

    elif tag_name == "table":
        process_table(element, doc)

    elif tag_name in ["div", "section", "article"]:
        # 容器元素，递归处理子元素
        for child in element.find_all(recursive=False):
            process_element(child, doc)

    elif tag_name == "br":
        # 换行
        doc.add_paragraph()

    elif tag_name in ["ul", "ol"]:
        process_list(element, doc)

    elif tag_name == "hr":
        # 水平线
        doc.add_paragraph("_" * 50)

    elif tag_name == "img":
        process_image(element, doc)

    else:
        # 默认处理：提取文本
        text = element.get_text().strip()
        if text and len(text) > 1:
            paragraph = doc.add_paragraph()
            add_text_to_paragraph(element, paragraph)


def process_heading(element, doc, level):
    """处理标题"""
    text = element.get_text().strip()
    if text:
        # 限制标题级别在 1-3
        actual_level = min(max(level, 1), 3)
        heading = doc.add_heading(level=actual_level)
        run = heading.add_run(text)

        # 设置标题样式
        if actual_level == 1:
            run.font.name = "黑体"
            run.font.size = Pt(16)
            run.font.bold = True
            try:
                run.element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
            except Exception:
                pass
        elif actual_level == 2:
            run.font.name = "黑体"
            run.font.size = Pt(14)
            run.font.bold = True
            try:
                run.element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
            except Exception:
                pass
        else:
            run.font.name = "宋体"
            run.font.size = Pt(12)
            run.font.bold = True
            try:
                run.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
            except Exception:
                pass


def process_paragraph(element, doc):
    """处理段落"""
    text = element.get_text().strip()
    if text or element.find_all(recursive=True):
        paragraph = doc.add_paragraph()
        add_text_to_paragraph(element, paragraph)

        # 设置段落对齐
        style = element.get("style", "")
        if "text-align: center" in style:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif "text-align: right" in style:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        elif "text-align: justify" in style:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def add_text_to_paragraph(element, paragraph):
    """向段落添加文本，处理内联样式"""
    for content in element.contents:
        if isinstance(content, str):
            # 普通文本
            text = content.strip()
            if text:
                run = paragraph.add_run(text)
                apply_default_font(run)
        elif hasattr(content, "name"):
            # HTML 元素
            tag_name = content.name.lower()

            if tag_name in ["strong", "b"]:
                # 粗体
                run = paragraph.add_run(content.get_text())
                run.bold = True
                apply_default_font(run)

            elif tag_name in ["em", "i"]:
                # 斜体
                run = paragraph.add_run(content.get_text())
                run.italic = True
                apply_default_font(run)

            elif tag_name == "u":
                # 下划线
                run = paragraph.add_run(content.get_text())
                run.underline = True
                apply_default_font(run)

            elif tag_name == "br":
                # 换行
                paragraph.add_run().add_break()

            elif tag_name == "span":
                # 内联样式
                run = paragraph.add_run(content.get_text())
                apply_span_styles(content, run)
                apply_default_font(run)

            else:
                # 递归处理嵌套元素
                add_text_to_paragraph(content, paragraph)


def apply_default_font(run):
    """应用默认字体"""
    run.font.name = "宋体"
    run.font.size = Pt(12)
    try:
        run.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    except Exception:
        pass


def apply_span_styles(span_element, run):
    """应用 span 元素的样式"""
    style = span_element.get("style", "")

    # 颜色
    color_match = re.search(r"color:\s*(#[0-9a-fA-F]+|\w+)", style)
    if color_match:
        color = color_match.group(1)
        if color.startswith("#"):
            try:
                # 转换十六进制颜色为 RGB
                hex_color = color.lstrip("#")
                rgb = tuple(int(hex_color[i : i + 2], 16) for i in (0, 2, 4))
                run.font.color.rgb = RGBColor(*rgb)
            except Exception:
                pass

    # 字体大小
    size_match = re.search(r"font-size:\s*(\d+)px", style)
    if size_match:
        try:
            size_pt = int(size_match.group(1)) * 0.75  # 近似转换 px 到 pt
            run.font.size = Pt(size_pt)
        except Exception:
            pass

    # 字体粗细
    if "font-weight: bold" in style or "font-weight: 700" in style:
        run.bold = True


def process_table(element, doc):
    """处理表格"""
    try:
        rows = element.find_all("tr")
        if not rows:
            return

        # 确定表格列数
        max_cols = 0
        for row in rows:
            cells = row.find_all(["td", "th"])
            max_cols = max(max_cols, len(cells))

        if max_cols == 0:
            return

        # 创建表格
        table = doc.add_table(rows=len(rows), cols=max_cols)
        table.style = "Table Grid"

        for i, row in enumerate(rows):
            cells = row.find_all(["td", "th"])
            for j, cell in enumerate(cells):
                if i < len(table.rows) and j < len(table.rows[i].cells):
                    cell_text = cell.get_text().strip()
                    table.cell(i, j).text = cell_text

                    # 设置表头样式
                    if cell.name == "th":
                        paragraph = table.cell(i, j).paragraphs[0]
                        run = paragraph.runs[0] if paragraph.runs else paragraph.add_run(cell_text)
                        run.bold = True

    except Exception as e:
        print(f"表格处理失败: {str(e)}")
        # 如果表格处理失败，添加文本内容
        table_text = element.get_text().strip()
        if table_text:
            doc.add_paragraph(f"[表格内容]: {table_text}")


def process_list(element, doc):
    """处理列表"""
    try:
        items = element.find_all("li")
        for item in items:
            text = item.get_text().strip()
            if text:
                if element.name == "ul":
                    # 无序列表
                    paragraph = doc.add_paragraph(style="List Bullet")
                else:
                    # 有序列表
                    paragraph = doc.add_paragraph(style="List Number")

                run = paragraph.add_run(text)
                apply_default_font(run)

    except Exception as e:
        print(f"列表处理失败: {str(e)}")


def process_image(element, doc):
    """处理图片"""
    try:
        src = element.get("src", "")
        if src and (src.startswith("http") or os.path.exists(src)):
            # 这里可以添加图片处理逻辑
            # 由于图片处理比较复杂，暂时跳过或添加占位符
            paragraph = doc.add_paragraph()
            run = paragraph.add_run("[图片]")
            run.italic = True
            run.font.color.rgb = RGBColor(128, 128, 128)
    except Exception as e:
        print(f"图片处理失败: {str(e)}")


def main():
    parser = argparse.ArgumentParser(description="HTML 转 Word 转换工具")
    parser.add_argument("-i", "--input", required=True, help="输入 HTML 文件路径")
    parser.add_argument("-o", "--output", required=True, help="输出 Word 文件路径")

    args = parser.parse_args()

    if not os.path.exists(args.input):
        print(f"错误: 输入文件不存在 - {args.input}")
        sys.exit(1)

    success = html_to_docx(args.input, args.output)
    sys.exit(0 if success else 1)


if __name__ == "__main__":
    main()
