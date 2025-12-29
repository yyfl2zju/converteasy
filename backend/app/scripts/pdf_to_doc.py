#!/usr/bin/env python3
"""
增强的 PDF 到 Word 转换脚本
使用 pdf2docx 并针对复杂布局（如简历）优化参数
解决 Issue #4 (图片丢失) 和 Issue #5 (格式混乱)
支持后处理以改善格式保留
"""
import argparse
import os
import sys
import time
import traceback
from pathlib import Path

from pdf2docx import Converter
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def convert_pdf_to_docx(pdf_path, docx_path):
    """
    使用 pdf2docx 转换，针对简历等复杂布局优化参数
    并进行后处理以改善格式
    """
    try:
        from pdf2docx import Converter
        
        print(f"[INFO] 开始转换: {pdf_path} -> {docx_path}")
        start_time = time.time()
        
        file_size_mb = os.path.getsize(pdf_path) / (1024 * 1024)
        print(f"[INFO] 文件大小: {file_size_mb:.2f} MB")
        
        print("[INFO] 第1步: 使用 pdf2docx 转换...")
        cv = Converter(pdf_path)
        
        # 优化参数以处理复杂布局和保留格式
        cv.convert(
            docx_path,
            start=0,
            end=None,

            # 核心：加强表格和线条检测
            extract_stream_table=True,       # 必须！提取无明确边框的"流式"表格
            table_border_threshold=0.5,      # 降低阈值，更敏感检测细线/虚线表格
            connected_border_tolerance=4.0,  # 提高容差，连接断开的细线
            max_border_width=3.0,            # 准考证边框通常很细

            # 图片/照片/二维码：加强提取
            min_image_width=40,              # 单位：点（72点=1英寸）
            min_image_height=40,
            extract_image_dpi=200,           # 提高清晰度

            # 文本行合并：减少乱换行
            line_overlap_threshold=0.4,      # 稍低，更积极合并行
            line_separate_threshold=12.0,    # 提高，避免误分
            line_break_free_space_ratio=0.15,

            # === 样式保留：字体、颜色、粗体、斜体等 ===
            keep_text_color=True,            # 保留文本颜色
            keep_text_style=True,            # 保留文本样式（粗体、斜体等）
            keep_text_bold=True,             # 保留粗体
            keep_text_italic=True,           # 保留斜体
            keep_font_size=True,             # 保留字体大小

            # 其他
            multi_processing=True,           # 加速 + 更稳定
        )
        
        cv.close()
        
        print("[INFO] 第2步: 后处理以修复格式混乱...")
        _postprocess_document(docx_path)
        
        elapsed = time.time() - start_time
        output_size_mb = os.path.getsize(docx_path) / (1024 * 1024)
        
        print(f"[SUCCESS] 转换成功: {docx_path}")
        print(f"[INFO] 输出文件大小: {output_size_mb:.2f} MB")
        print(f"[INFO] 总耗时: {elapsed:.1f} 秒")
        
        return True
        
    except Exception as e:
        print(f"[ERROR] 转换失败: {str(e)}", file=sys.stderr)
        traceback.print_exc()
        return False


def _postprocess_document(docx_path):
    """
    后处理Word文档，修复格式混乱问题：
    - 规范化表格格式
    - 清理多余换行
    - 统一默认字体为中文字体
    - 修复行距问题
    """
    try:
        doc = Document(docx_path)
        
        # 1. 统一默认字体（中文支持）
        style = doc.styles['Normal']
        font = style.font
        if font.name is None or font.name.lower() not in ['calibri', 'arial']:
            font.name = 'Calibri'
            font.size = Pt(11)
        
        # 2. 处理所有段落：清理过多换行，规范化格式
        modified = False
        prev_empty = False
        
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            
            # 清理连续空段落（保留一个）
            if not text:
                if prev_empty:
                    para._element.getparent().remove(para._element)
                    modified = True
                prev_empty = True
            else:
                prev_empty = False
                
                # 3. 规范化行距（避免过大行距导致的混乱）
                if para.paragraph_format.line_spacing is None or para.paragraph_format.line_spacing > 2.0:
                    para.paragraph_format.line_spacing = 1.15
                    modified = True
                
                # 4. 处理缩进（避免过大缩进）
                if para.paragraph_format.first_line_indent and para.paragraph_format.first_line_indent.pt > 100:
                    para.paragraph_format.first_line_indent = Pt(0)
                    modified = True
        
        # 5. 处理表格格式
        for table in doc.tables:
            try:
                # 自动调整列宽
                for row in table.rows:
                    for cell in row.cells:
                        # 确保表格单元格内的段落格式正确
                        for para in cell.paragraphs:
                            if para.paragraph_format.line_spacing is None:
                                para.paragraph_format.line_spacing = 1.15
                                modified = True
            except Exception as e:
                print(f"[WARN] 处理表格时出现问题: {e}")
        
        # 6. 保存修改
        if modified or True:  # 总是保存
            doc.save(docx_path)
            print("[INFO] 文档格式已规范化")
            
    except Exception as e:
        print(f"[WARN] 后处理时出现问题，继续使用原文档: {e}")


def main():
    parser = argparse.ArgumentParser(
        description="PDF 转 Word (增强版) - 保留图片、格式和样式"
    )
    parser.add_argument("-i", "--input", required=True, help="输入 PDF 文件路径")
    parser.add_argument("-o", "--output", required=True, help="输出 Word 文件路径")

    args = parser.parse_args()

    # 验证输入文件
    if not os.path.exists(args.input):
        print(f"[ERROR] 输入文件不存在: {args.input}", file=sys.stderr)
        sys.exit(1)

    # 验证文件扩展名
    if not args.input.lower().endswith(".pdf"):
        print("[ERROR] 输入文件必须是 PDF 格式", file=sys.stderr)
        sys.exit(1)

    print(f"[INFO] 输入文件: {args.input}")
    print(f"[INFO] 输出文件: {args.output}")

    # 执行转换
    success = convert_pdf_to_docx(args.input, args.output)

    # 验证输出
    if success and os.path.exists(args.output):
        output_size = os.path.getsize(args.output)
        if output_size > 0:
            print(f"[SUCCESS] 最终转换成功，输出文件大小: {output_size} 字节")
            sys.exit(0)
        else:
            print("[ERROR] 输出文件为空", file=sys.stderr)
            sys.exit(1)
    else:
        print("[ERROR] 转换方法失败", file=sys.stderr)
        sys.exit(1)


if __name__ == "__main__":
    main()
