#!/usr/bin/env python3
import argparse
import pandas as pd
from docx import Document
import sys
import os


def xls_to_doc(xls_path, doc_path):
    """将 Excel 文件转换为 Word 文档"""
    try:
        # 读取 Excel 文件
        df = pd.read_excel(xls_path)

        # 创建 Word 文档
        doc = Document()

        # 添加标题
        doc.add_heading("Excel 数据", level=1)

        # 添加表格到 Word
        table = doc.add_table(rows=len(df) + 1, cols=len(df.columns))

        # 添加表头
        for col_idx, column_name in enumerate(df.columns):
            table.cell(0, col_idx).text = str(column_name)

        # 添加数据行
        for row_idx, row in df.iterrows():
            for col_idx, value in enumerate(row):
                table.cell(row_idx + 1, col_idx).text = str(value)

        doc.save(doc_path)
        print(f"转换成功: {xls_path} -> {doc_path}")
        return True
    except Exception as e:
        print(f"转换失败: {str(e)}")
        return False


def main():
    parser = argparse.ArgumentParser(description="Excel 转 Word")
    parser.add_argument("-i", "--input", required=True, help="输入 Excel 文件路径")
    parser.add_argument("-o", "--output", required=True, help="输出 Word 文件路径")

    args = parser.parse_args()

    if not os.path.exists(args.input):
        print(f"错误: 输入文件不存在 {args.input}")
        sys.exit(1)

    success = xls_to_doc(args.input, args.output)
    sys.exit(0 if success else 1)


if __name__ == "__main__":
    main()
