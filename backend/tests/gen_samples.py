"""
Generate small sample files for conversion testing.
Outputs to `backend/tests/samples/`.
"""
from pathlib import Path

SAMPLES_DIR = Path(__file__).parent / "samples"
SAMPLES_DIR.mkdir(parents=True, exist_ok=True)


def make_txt():
    (SAMPLES_DIR / "sample.txt").write_text("Hello ConvertEasy!\n这是一个测试文本。\n", encoding="utf-8")


def make_html():
    (SAMPLES_DIR / "sample.html").write_text(
        """<!doctype html><html><head><meta charset='utf-8'><title>Sample</title></head>
        <body><h1>ConvertEasy</h1><p>Simple HTML sample.</p></body></html>""",
        encoding="utf-8",
    )


def make_docx():
    try:
        from docx import Document
    except Exception:
        print("python-docx not installed; skip DOCX")
        return
    doc = Document()
    doc.add_heading("ConvertEasy Sample", level=1)
    doc.add_paragraph("This is a small DOCX file used for testing.")
    doc.save(SAMPLES_DIR / "sample.docx")


def make_xlsx():
    try:
        from openpyxl import Workbook
    except Exception:
        print("openpyxl not installed; skip XLSX")
        return
    wb = Workbook()
    ws = wb.active
    ws.title = "Sheet1"
    ws["A1"] = "ConvertEasy"
    ws["B1"] = "Sample"
    ws["A2"] = 123
    ws["B2"] = 456
    wb.save(SAMPLES_DIR / "sample.xlsx")


def make_wav(duration_sec: float = 0.25, sample_rate: int = 8000):
    # Generate a tiny silent WAV using stdlib
    import wave
    import struct

    frames = int(duration_sec * sample_rate)
    with wave.open(str(SAMPLES_DIR / "sample.wav"), "w") as w:
        w.setnchannels(1)
        w.setsampwidth(2)  # 16-bit PCM
        w.setframerate(sample_rate)
        silence = struct.pack("<h", 0)
        for _ in range(frames):
            w.writeframes(silence)


def main():
    make_txt()
    make_html()
    make_docx()
    make_xlsx()
    make_wav()
    print(f"Samples written to: {SAMPLES_DIR}")


if __name__ == "__main__":
    main()
